#!/usr/bin/env python3
"""
convert_md_to_docx.py - Convert markdown to .docx files

Usage:
    python convert_md_to_docx.py <input.md>

Output:
    <input>.docx

Features:
    - Bold (**text**) -> Bold formatting
    - Italic (*text*) -> Italic formatting
    - Inline code (`text`) -> Monospace font
    - Code blocks (```) -> Bordered box with code font
    - --- -> Page break
    - Headings -> Document headings (H1-H4)
    - Lists -> Bullet/numbered lists
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn


def parse_markdown(content):
    """Parse markdown into structured sections."""
    sections = []
    current_section = {"level": 0, "title": "", "content": [], "items": [], "code_blocks": [], "has_page_break": False}
    
    in_code_block = False
    code_lines = []
    code_lang = ""
    
    for line in content.split('\n'):
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                if code_lines:
                    current_section["code_blocks"].append({
                        "lang": code_lang,
                        "code": '\n'.join(code_lines)
                    })
                code_lines = []
                code_lang = ""
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                code_lang = line.strip()[3:].strip()
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Detect horizontal rules (page breaks)
        if re.match(r'^---+\s*$', line):
            current_section["has_page_break"] = True
            continue
        # Detect headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            if current_section["title"] or current_section["content"] or current_section["code_blocks"]:
                sections.append(current_section)
            current_section = {
                "level": len(heading_match.group(1)),
                "title": heading_match.group(2).strip(),
                "content": [],
                "items": [],
                "code_blocks": [],
                "has_page_break": False
            }
        # Detect list items
        elif re.match(r'^\s*[-*]\s+', line):
            item = re.sub(r'^\s*[-*]\s+', '', line).strip()
            current_section["items"].append(item)
        # Detect numbered items
        elif re.match(r'^\s*\d+\.\s+', line):
            item = re.sub(r'^\s*\d+\.\s+', '', line).strip()
            current_section["items"].append(item)
        # Regular content
        elif line.strip():
            current_section["content"].append(line.strip())
    
    if current_section["title"] or current_section["content"] or current_section["code_blocks"]:
        sections.append(current_section)
    
    return sections


def add_formatted_paragraph(paragraph, text):
    """Add text with bold, italic, and inline code formatting."""
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            if plain_text:
                paragraph.add_run(plain_text)
        
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        
        last_end = match.end()
    
    if last_end < len(text):
        plain_text = text[last_end:]
        if plain_text:
            paragraph.add_run(plain_text)


def add_code_block(doc, code_text, lang=""):
    """Add a code block with border and monospace font."""
    p = doc.add_paragraph()
    
    # Set paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    for side in ['top', 'left', 'bottom', 'right']:
        bdr = pBdr.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '4',
            qn('w:color'): '999999',
        })
        pBdr.append(bdr)
    pPr.append(pBdr)
    
    # Set paragraph shading (light gray background)
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5',
    })
    pPr.append(shd)
    
    # Add language label if present
    if lang:
        run = p.add_run(f"{lang}\n")
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
    
    # Add code content
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # Set left indent to give some padding
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def create_docx(sections, output_path):
    """Create a Word document from parsed sections."""
    doc = Document()
    
    # Add title
    if sections and sections[0]["title"]:
        doc.add_heading(sections[0]["title"], level=0)
    
    for section in sections[1:]:
        if not section["title"]:
            continue
            
        # Add page break if needed
        if section["has_page_break"]:
            doc.add_page_break()
        
        # Add heading
        doc.add_heading(section["title"], level=min(section["level"], 4))
        
        # Add content paragraphs
        for para in section["content"]:
            if para:
                p = doc.add_paragraph()
                add_formatted_paragraph(p, para)
        
        # Add code blocks
        for code_block in section["code_blocks"]:
            add_code_block(doc, code_block["code"], code_block["lang"])
        
        # Add list items
        for item in section["items"]:
            if item:
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_paragraph(p, item)
    
    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_md_to_docx.py <input.md>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_dir = os.path.dirname(input_path) if os.path.dirname(input_path) else "."
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    sections = parse_markdown(content)
    
    docx_path = os.path.join(output_dir, base_name + ".docx")
    create_docx(sections, docx_path)


if __name__ == '__main__':
    main()
